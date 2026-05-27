import re
import io
import os
import json
import html as html_mod
from datetime import datetime
from difflib import SequenceMatcher

import anthropic
import pdfplumber
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv

load_dotenv()

st.set_page_config(
    page_title="Resume ATS Optimizer",
    page_icon="📄",
    layout="wide",
)

# ── System prompt ─────────────────────────────────────────────────────────────────
SYSTEM_PROMPT = """You are a world-class ATS (Applicant Tracking System) resume optimizer
and career strategist. You have deep expertise spanning over a decade of helping candidates
land interviews at top companies by making their resumes both ATS-compatible and compelling
to human reviewers.

## Your Core Expertise

**ATS Systems**: You have comprehensive knowledge of modern applicant tracking systems —
Workday, Greenhouse, Lever, iCIMS, Taleo, BambooHR, Jobvite, SmartRecruiters, and others.
You understand exactly how they tokenize and parse resume text, how they score keyword
matches (exact, partial, and semantic), how they recognize section headers, how they rank
candidates against each other, and what causes automatic rejection before a human ever sees
the file.

**Recruiting and HR Practices**: You understand recruiter workflows end-to-end — how they
construct Boolean search strings to surface candidates, how they screen 200 resumes in two
hours (typically under 10 seconds per resume), which visual patterns catch attention, and
which patterns trigger immediate passes. You know what "minimum viable" looks like for each
level of seniority and each industry.

**Industry Terminology**: You are fluent in the vocabulary, frameworks, tools, certifications,
and career paths across technology (software engineering, data science, ML/AI, DevOps, cloud
infrastructure, product management, cybersecurity, QA), finance (investment banking, fintech,
asset management, accounting, risk and compliance), healthcare (clinical operations, biotech,
health IT, medical devices), marketing (growth, brand, content, performance, demand
generation), operations (supply chain, project management, consulting, strategy), legal, and
more. You understand seniority-specific expectations at each level from entry to executive.

**Resume Writing**: You know how to write quantified achievement bullets using the STAR and
XYZ frameworks, how to craft a tailored professional summary, how to structure a skills
section for maximum keyword coverage without looking spammy, how to use section headers that
every ATS reliably recognizes (Experience, Education, Skills, Summary, Certifications), and
how to format a document so it parses cleanly through automated systems and still impresses
human reviewers.

## Guiding Principles

When analyzing and optimizing resumes, you strictly follow these principles:

1. **Integrity first**: You never suggest fabricating, exaggerating, or misrepresenting
   experience, education, titles, dates, or metrics. You only help candidates express
   their genuine background in the strongest, most industry-standard terms possible.

2. **Mirror the employer's language**: You identify the exact verbs, nouns, and phrases the
   employer uses in the job description and recommend weaving them into the resume wherever
   honest and accurate — because ATS systems often do exact-string matching before any
   semantic analysis.

3. **Quantify everything possible**: Vague bullets ("managed projects", "worked with clients")
   score poorly with both ATS and humans. You transform them into specific, metrics-driven
   statements that demonstrate scope, impact, and ownership.

4. **Prioritize ruthlessly**: Not every gap matters equally. You rank suggestions by expected
   ATS score impact and candidate effort required, and you always highlight the three to five
   changes most likely to meaningfully move the needle.

5. **ATS-safe structure**: You flag formatting hazards — tables, multi-column layouts, text
   boxes, headers/footers with critical content, non-standard bullet characters, images or
   logos, and decorative horizontal rules — and recommend plain-text alternatives.

6. **Strategic keyword placement**: High-priority keywords should appear in the professional
   summary (highest ATS weight), job titles and bullet points (high weight), and the skills
   section (medium weight). Appearing in only one section is weaker than appearing in three.

7. **Conciseness for one-page fit**: All suggested content must be concise. Rewritten bullets
   must be under 120 characters. The goal is a final resume that fits on one letter-size page.

8. **No em dashes**: Never use em dashes (—) in any rewritten resume content (summary, skills,
   bullets). Use commas, periods, or parentheses instead. This applies to all user-facing
   resume text — the JSON field values themselves must not contain "—".

## Response Format

You always respond with a single, valid JSON object using exactly this schema. No markdown,
no code fences, no explanatory text before or after the JSON:

{
  "ats_score_estimate": <integer 0-100, honest estimate of current ATS compatibility>,
  "score_breakdown": {
    "keyword_match":    <integer 0-100>,
    "quantification":   <integer 0-100>,
    "formatting":       <integer 0-100>,
    "relevance":        <integer 0-100>
  },
  "keyword_gaps": {
    "technical_skills":   ["missing technical skills mentioned or implied in JD"],
    "soft_skills":        ["missing soft skills and interpersonal competencies"],
    "tools_technologies": ["missing tools, software, platforms, frameworks"],
    "certifications":     ["relevant certifications not present in resume"],
    "domain_knowledge":   ["industry concepts, methodologies, frameworks not mentioned"]
  },
  "strengths": [
    "2–5 specific strengths already well-represented that align with this role"
  ],
  "weaknesses": [
    "2–4 specific weaknesses or gaps that hurt this resume's chances for this specific role"
  ],
  "current_summary": "The candidate's existing professional summary verbatim, or empty string if none found",
  "current_skills": "The candidate's existing skills section verbatim, or empty string if none found",
  "suggested_skills": "Complete rewritten skills section integrating the most critical missing keywords naturally, formatted to match the original style (comma-separated, categorized, etc.)",
  "critical_suggestions": [
    "3–5 highest-priority, actionable improvements ranked by expected ATS impact (one sentence each)"
  ],
  "bullet_rewrites": [
    {
      "original":  "exact bullet or phrase from the resume",
      "suggested": "rewritten version — strong verb, quantified result, under 120 characters",
      "reason":    "one sentence explaining the ATS or recruiter impact of this change"
    }
  ],
  "keywords_to_add_to_skills": [
    "keywords and phrases to add explicitly to the skills section"
  ],
  "optimized_summary": "2–3 sentence professional summary for this role, under 400 characters total",
  "general_advice": [
    "3–5 concise strategic advice bullets specific to this candidate and role"
  ]
}"""


# ── Constants ─────────────────────────────────────────────────────────────────────
BULLET_CHARS = frozenset("•·◦▪▸▹▶►●○◆◇✓✔➢➤→–—")

_DATE_SUFFIX = re.compile(
    r"(?:\t|\s{2,})"
    r"("
    r"(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\.?\s+)?"
    r"\d{4}"
    r"(?:\s*[–—\-]+\s*"
    r"(?:(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\.?\s+)?"
    r"\d{4}|Present|Current|Now))?"
    r")\s*$",
    re.IGNORECASE,
)


# ── Extraction ────────────────────────────────────────────────────────────────────

def extract_pdf(data: bytes) -> str:
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]
    return "\n".join(pages).strip()


def extract_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    lines = []
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        style_name = (p.style.name or "").lower()
        is_list = "list" in style_name or p._p.find(qn("w:numPr")) is not None
        parts = []
        for run in p.runs:
            if not run.text:
                continue
            text = run.text
            if run.bold and text.strip():
                inner = text.strip()
                lead = text[: len(text) - len(text.lstrip())]
                trail = text[len(text.rstrip()):]
                parts.append(f"{lead}**{inner}**{trail}")
            else:
                parts.append(text)
        line = "".join(parts).strip()
        if not line:
            continue
        if is_list:
            line = f"• {line}"
        lines.append(line)
    return "\n".join(lines)


# ── API calls ─────────────────────────────────────────────────────────────────────

def _get_client() -> anthropic.Anthropic:
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError(
            "ANTHROPIC_API_KEY is not set. Add it to your .env file and restart the app."
        )
    return anthropic.Anthropic(api_key=api_key)


def analyze_resume(resume_text: str, jd: str) -> dict:
    response = _get_client().messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4096,
        system=[{"type": "text", "text": SYSTEM_PROMPT, "cache_control": {"type": "ephemeral"}}],
        messages=[
            {
                "role": "user",
                "content": (
                    "Analyze the following resume against the job description. "
                    "Reply with only the JSON object, no other text.\n\n"
                    f"## RESUME\n\n{resume_text}\n\n"
                    f"## JOB DESCRIPTION\n\n{jd}"
                ),
            }
        ],
    )
    raw = response.content[0].text.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw.strip())


def auto_format_resume(text: str) -> str:
    """Restructure resume text for optimal ATS formatting without changing content."""
    response = _get_client().messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4096,
        messages=[
            {
                "role": "user",
                "content": (
                    "Reformat the resume below for maximum ATS formatting score. "
                    "Preserve every fact, number, and accomplishment verbatim. Only restructure.\n\n"
                    "RULES:\n"
                    "- Line 1: candidate's full name (centered)\n"
                    "- Lines 2-3: contact info (email, phone, city/state, LinkedIn) separated by ' | '\n"
                    "- Use these standard ALL-CAPS section headers exactly: SUMMARY, EXPERIENCE, "
                    "EDUCATION, SKILLS, PROJECTS, CERTIFICATIONS (only the ones the resume has)\n"
                    "- Each job entry: bold company and title on one line; dates on the same line, "
                    "separated from the title by 2+ spaces (will render right-aligned)\n"
                    "- All bullet points start with '• ' (bullet + space). One sentence each\n"
                    "- Use **bold** for company names, job titles, and degree names\n"
                    "- Single column, no tables, no special characters or emojis\n"
                    "- Do NOT use em dashes (—). Use hyphens, commas, or periods instead\n"
                    "- Return ONLY the reformatted resume text, no commentary\n\n"
                    f"{text}"
                ),
            }
        ],
    )
    return response.content[0].text.strip()


def generate_cover_letter(resume_text: str, jd: str) -> str:
    today = datetime.now().strftime("%B %d, %Y")
    response = _get_client().messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1024,
        messages=[
            {
                "role": "user",
                "content": (
                    f"Today's date is {today}.\n\n"
                    "Write a professional cover letter tailored to the job description below.\n\n"
                    "STRICT FORMAT (one line per item, exactly):\n"
                    "Line 1: Full name (from resume)\n"
                    "Line 2: City, State (from resume)\n"
                    "Line 3: Email address (from resume)\n"
                    "Line 4: Phone number (from resume)\n"
                    "Line 5: LinkedIn URL (from resume — omit this line entirely if not present)\n"
                    "[blank line]\n"
                    f"{today}\n"
                    "[blank line]\n"
                    "Dear Hiring Manager, (or specific name if mentioned in the JD)\n"
                    "[blank line]\n"
                    "3 short body paragraphs separated by blank lines:\n"
                    "  - Opening (50-70 words): specific enthusiasm for this role and company\n"
                    "  - Body (90-120 words): 2 quantified achievements from the resume that match the JD\n"
                    "  - Closing (40-60 words): brief, express eagerness to discuss\n"
                    "[blank line]\n"
                    "Sincerely,\n"
                    "[Full name from resume]\n\n"
                    "REQUIREMENTS:\n"
                    "- Body paragraphs total 200–250 words (must fit one letter-size page with header)\n"
                    "- Tone: confident and direct, no clichés like 'I am writing to apply'\n"
                    "- Naturally weave keywords from the JD\n"
                    "- Do NOT use em dashes (—). Use commas, periods, or hyphens instead\n"
                    "- Return ONLY the letter text starting with the name, nothing else\n\n"
                    f"## RESUME\n\n{resume_text}\n\n"
                    f"## JOB DESCRIPTION\n\n{jd}"
                ),
            }
        ],
    )
    return response.content[0].text.strip()


# ── Formatting primitives ─────────────────────────────────────────────────────────

def _is_section_header(s: str) -> bool:
    clean = re.sub(r"\*\*", "", s)
    return (
        clean.isupper()
        and bool(re.match(r"^[A-Z][A-Z\s&/]+$", clean))
        and 5 < len(clean) < 42
    )


def _is_bullet(s: str) -> bool:
    if not s:
        return False
    if s[0] in BULLET_CHARS and len(s) > 1:
        return True
    if len(s) > 2 and s[0] in "-*" and s[1] == " ":
        return True
    return False


def _bullet_content(s: str) -> str:
    if s and s[0] in BULLET_CHARS:
        return s[1:].lstrip()
    if len(s) > 1 and s[0] in "-*" and s[1] == " ":
        return s[2:]
    return s


def _split_date_line(line: str):
    m = _DATE_SUFFIX.search(line)
    if m and m.start() > 0:
        return line[: m.start()].rstrip(), m.group(1).strip()
    return line, None


def _render_inline(raw: str) -> str:
    escaped = html_mod.escape(raw)
    return re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", escaped)


def _add_bold_runs(para, text: str, size: Pt, font: str = "Arial") -> None:
    for part in re.split(r"(\*\*.+?\*\*)", text):
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            run = para.add_run(part)
        run.font.name = font
        run.font.size = size


def _add_bottom_border(para) -> None:
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "333333")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_right_tab(para, twips: int = 10368) -> None:
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(twips))
    tabs.append(tab)
    pPr.append(tabs)


# ── Word export: resume ───────────────────────────────────────────────────────────

def to_word(text: str) -> bytes:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.5)
        sec.left_margin = sec.right_margin = Inches(0.65)

    lines = text.split("\n")
    first_nonempty = next((i for i, ln in enumerate(lines) if ln.strip()), 0)

    for idx, line in enumerate(lines):
        s = line.strip()

        if not s:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            continue

        if idx == first_nonempty:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            _add_bold_runs(p, s, Pt(15))

        elif idx - first_nonempty <= 3 and re.search(r"[@|]|\d{3}", s):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            _add_bold_runs(p, s, Pt(9))

        elif _is_section_header(s):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            clean = re.sub(r"\*\*", "", s)
            run = p.add_run(clean)
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.bold = True
            _add_bottom_border(p)

        elif _is_bullet(s):
            content = _bullet_content(s)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            _add_bold_runs(p, f"•  {content}", Pt(9.5))

        else:
            main, date = _split_date_line(s)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            if date:
                _add_right_tab(p)
                _add_bold_runs(p, main, Pt(9.5))
                run = p.add_run(f"\t{date}")
                run.font.name = "Arial"
                run.font.size = Pt(9.5)
            else:
                _add_bold_runs(p, s, Pt(9.5))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Word export: cover letter ─────────────────────────────────────────────────────

def to_word_letter(text: str) -> bytes:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.8)
        sec.left_margin = sec.right_margin = Inches(1)

    for line in text.split("\n"):
        s = line.strip()
        p = doc.add_paragraph()
        if s:
            run = p.add_run(s)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
        else:
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.0

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── HTML previews ─────────────────────────────────────────────────────────────────

def format_resume_html(text: str) -> str:
    lines = text.split("\n")
    first_nonempty = next((i for i, ln in enumerate(lines) if ln.strip()), 0)

    parts = [
        "<div style='font-family:Arial,sans-serif;width:100%;padding:18px 22px;"
        "background:#fff;border:1px solid #d0d0d0;border-radius:3px;"
        "line-height:1.38;color:#111;box-sizing:border-box;font-size:9.5pt;'>"
    ]

    for i, line in enumerate(lines):
        raw = line.strip()
        if not raw:
            parts.append("<div style='height:4px'></div>")
            continue

        if i == first_nonempty:
            parts.append(
                f"<div style='text-align:center;font-size:15pt;font-weight:700;"
                f"margin-bottom:2px'>{_render_inline(raw)}</div>"
            )
        elif i - first_nonempty <= 3 and re.search(r"[@|]|\d{3}", raw):
            parts.append(
                f"<div style='text-align:center;font-size:8.5pt;color:#555;"
                f"margin-bottom:2px'>{_render_inline(raw)}</div>"
            )
        elif _is_section_header(raw):
            clean = html_mod.escape(re.sub(r"\*\*", "", raw))
            parts.append(
                f"<div style='font-weight:700;font-size:9.5pt;letter-spacing:.6px;"
                f"border-bottom:1.2px solid #222;padding-bottom:1px;"
                f"margin-top:10px;margin-bottom:3px'>{clean}</div>"
            )
        elif _is_bullet(raw):
            content = _render_inline(_bullet_content(raw))
            parts.append(
                f"<div style='margin-left:14px;margin-bottom:1px'>"
                f"&bull;&nbsp;{content}</div>"
            )
        else:
            main, date = _split_date_line(raw)
            if date:
                parts.append(
                    f"<div style='display:flex;justify-content:space-between;"
                    f"margin-bottom:1px'>"
                    f"<span>{_render_inline(main)}</span>"
                    f"<span style='white-space:nowrap;padding-left:8px'>"
                    f"{html_mod.escape(date)}</span></div>"
                )
            else:
                parts.append(f"<div style='margin-bottom:1px'>{_render_inline(raw)}</div>")

    parts.append("</div>")
    return "\n".join(parts)


def format_letter_html(text: str) -> str:
    parts = [
        "<div style='font-family:Arial,sans-serif;width:100%;padding:24px 28px;"
        "background:#fff;border:1px solid #d0d0d0;border-radius:3px;"
        "line-height:1.35;color:#111;box-sizing:border-box;font-size:11pt;'>"
    ]
    for line in text.split("\n"):
        raw = line.strip()
        if raw:
            parts.append(f"<div style='margin-bottom:1px'>{html_mod.escape(raw)}</div>")
        else:
            parts.append("<div style='height:9px'></div>")
    parts.append("</div>")
    return "\n".join(parts)


# ── Misc helpers ──────────────────────────────────────────────────────────────────

def _fuzzy_replace(text: str, original: str, replacement: str, threshold: float = 0.72):
    """Replace `original` in `text`, with multi-line block fallback for sections."""
    if original in text:
        return text.replace(original, replacement, 1), True

    norm_orig = " ".join(original.split())
    n_orig_lines = max(1, len([ln for ln in original.split("\n") if ln.strip()]))
    lines = text.split("\n")
    best_ratio, best_start, best_end = 0.0, -1, -1

    max_window = min(len(lines), max(n_orig_lines + 3, 6))
    for window in range(1, max_window + 1):
        for i in range(len(lines) - window + 1):
            block = "\n".join(lines[i : i + window])
            if not block.strip():
                continue
            clean = re.sub(r"\*\*", "", block)
            ratio = SequenceMatcher(None, norm_orig, " ".join(clean.split())).ratio()
            if ratio > best_ratio:
                best_ratio, best_start, best_end = ratio, i, i + window

    if best_ratio >= threshold and best_start >= 0:
        new_lines = lines[:best_start] + replacement.split("\n") + lines[best_end:]
        return "\n".join(new_lines), True

    return text, False


def priority_badge(i: int) -> str:
    if i <= 1:
        return "🔴 High"
    if i <= 3:
        return "🟡 Medium"
    return "🟢 Low"


def _clear_rewrite_keys() -> None:
    for k in list(st.session_state.keys()):
        if k.startswith("rw_edit_") or k in ("summary_edit", "skills_edit", "resume_editor"):
            del st.session_state[k]


# ── Session state ─────────────────────────────────────────────────────────────────
defaults = {
    "analysis": None,
    "resume_text": "",
    "export_filename": "optimized_resume.docx",
    "last_uploaded": None,
    "rescore": None,
    "analysis_jd": "",
    "cover_letter": "",
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ── Header ────────────────────────────────────────────────────────────────────────
st.title("📄 Resume ATS Optimizer")
st.caption(
    "Upload your resume and paste the job description to receive keyword gap analysis, "
    "ATS-optimized suggestions, and a tailored cover letter."
)
st.divider()

# ── Inputs ────────────────────────────────────────────────────────────────────────
col_resume, col_jd = st.columns(2)

with col_resume:
    st.subheader("① Upload Your Resume")
    uploaded = st.file_uploader("PDF or DOCX", type=["pdf", "docx"])

    if uploaded:
        if uploaded.name != st.session_state.last_uploaded:
            raw_bytes = uploaded.getvalue()
            try:
                if uploaded.name.lower().endswith(".pdf"):
                    extracted = extract_pdf(raw_bytes)
                else:
                    extracted = extract_docx(raw_bytes)

                if extracted:
                    st.session_state.resume_text = extracted
                    st.session_state.last_uploaded = uploaded.name
                    st.session_state.export_filename = (
                        uploaded.name.rsplit(".", 1)[0] + "_optimized.docx"
                    )
                    st.session_state.analysis = None
                    st.session_state.rescore = None
                    st.session_state.cover_letter = ""
                    _clear_rewrite_keys()
                    for k in ("cl_editor",):
                        if k in st.session_state:
                            del st.session_state[k]
                else:
                    st.error(
                        "No text found. The file may be image-based or password-protected."
                    )
            except Exception as exc:
                st.error(f"Could not read file: {exc}")

        if st.session_state.resume_text:
            word_count = len(st.session_state.resume_text.split())
            st.success(f"✅ {word_count:,} words loaded from **{uploaded.name}**")

with col_jd:
    st.subheader("② Paste Job Description")
    jd = st.text_area(
        "Full job posting",
        height=280,
        placeholder=(
            "Paste the complete job description here — include the responsibilities, "
            "required qualifications, and preferred qualifications sections…"
        ),
        label_visibility="collapsed",
    )

st.divider()

# ── Analyze button ────────────────────────────────────────────────────────────────
ready = bool(st.session_state.resume_text and jd.strip())
if st.button("🔍 Analyze Resume", type="primary", disabled=not ready):
    _clear_rewrite_keys()
    with st.spinner("Analyzing your resume against the job description…"):
        try:
            st.session_state.analysis = analyze_resume(st.session_state.resume_text, jd)
            st.session_state.analysis_jd = jd
            st.session_state.rescore = None
        except json.JSONDecodeError:
            st.error("The model returned an unexpected format. Please try again.")
        except RuntimeError as exc:
            st.error(str(exc))
        except anthropic.APIError as exc:
            st.error(f"Anthropic API error: {exc}")

if not ready and not st.session_state.resume_text:
    st.info("Upload a resume and paste a job description to get started.")

# ── Results ───────────────────────────────────────────────────────────────────────
if st.session_state.analysis:
    a = st.session_state.analysis

    # 1 ── Scores ─────────────────────────────────────────────────────────────────
    st.subheader("📊 Analysis Results")
    score = int(a.get("ats_score_estimate", 0))
    bk = a.get("score_breakdown", {})
    m1, m2, m3, m4, m5 = st.columns(5)
    m1.metric("Overall ATS Score", f"{score} / 100")
    m2.metric("Keyword Match", bk.get("keyword_match", "—"))
    m3.metric("Quantification", bk.get("quantification", "—"))
    m4.metric("Formatting", bk.get("formatting", "—"))
    m5.metric("Relevance", bk.get("relevance", "—"))
    st.progress(min(score, 100) / 100)

    st.divider()

    # 2 ── Strengths | Weaknesses ─────────────────────────────────────────────────
    sw_col, wk_col = st.columns(2)
    with sw_col:
        strengths = a.get("strengths", [])
        if strengths:
            st.subheader("✅ Strengths")
            for s in strengths:
                st.success(s)
    with wk_col:
        weaknesses = a.get("weaknesses", [])
        if weaknesses:
            st.subheader("⚠️ Weaknesses")
            for w in weaknesses:
                st.warning(w)

    # 3 ── Strategic Advice ───────────────────────────────────────────────────────
    advice = a.get("general_advice", [])
    if advice:
        st.subheader("💼 Strategic Advice")
        if isinstance(advice, list):
            for pt in advice:
                st.markdown(f"- {pt}")
        else:
            for sent in re.split(r"\.\s+", str(advice)):
                if sent.strip():
                    st.markdown(f"- {sent.strip().rstrip('.')}.")

    st.divider()

    # 4 ── Suggested Rewrites ─────────────────────────────────────────────────────
    st.subheader("📋 Suggested Rewrites")
    st.caption(
        "Edit any **After** field, then click Apply to push all changes into the preview. "
        "Use `**word**` in Edit Source to bold specific words."
    )

    opt_summary = a.get("optimized_summary", "")
    rewrites = a.get("bullet_rewrites", [])

    if opt_summary:
        current_summary = a.get("current_summary", "")
        snip = (current_summary[:55] + "…") if len(current_summary) > 55 else current_summary
        exp_label = f'📝 Professional Summary — "{snip}"' if snip else "📝 Professional Summary"
        with st.expander(exp_label, expanded=True):
            st.markdown("**Priority: 🔴 High**")
            bc, ac = st.columns(2)
            with bc:
                st.markdown("**Before**")
                st.warning(current_summary if current_summary else "_No existing summary detected_")
            with ac:
                st.markdown("**After (editable)**")
                if "summary_edit" not in st.session_state:
                    st.session_state["summary_edit"] = opt_summary
                st.text_area("Optimized summary", key="summary_edit", height=110,
                             label_visibility="collapsed")
            st.caption("💡 Leads with high-priority keywords and mirrors the role's language.")

    suggested_skills = a.get("suggested_skills", "")
    if suggested_skills:
        current_skills = a.get("current_skills", "")
        with st.expander("🛠️ Skills Section", expanded=True):
            st.markdown("**Priority: 🔴 High**")
            bc, ac = st.columns(2)
            with bc:
                st.markdown("**Before**")
                st.warning(current_skills if current_skills else "_No existing skills section detected_")
            with ac:
                st.markdown("**After (editable)**")
                if "skills_edit" not in st.session_state:
                    st.session_state["skills_edit"] = suggested_skills
                st.text_area("Suggested skills", key="skills_edit", height=110,
                             label_visibility="collapsed")
            st.caption("💡 Integrates missing keywords while preserving your existing skills format.")

    for i, rw in enumerate(rewrites):
        orig = rw.get("original", "")
        snip = f'"{orig[:55]}…"' if len(orig) > 55 else f'"{orig}"'
        key = f"rw_edit_{i}"
        with st.expander(f"✏️ Rewrite #{i + 1} — {snip}"):
            st.markdown(f"**Priority: {priority_badge(i)}**")
            bc, ac = st.columns(2)
            with bc:
                st.markdown("**Before**")
                st.warning(orig)
            with ac:
                st.markdown("**After (editable)**")
                if key not in st.session_state:
                    st.session_state[key] = rw.get("suggested", "")
                st.text_area("Suggested rewrite", key=key, height=110,
                             label_visibility="collapsed")
            st.caption(f"💡 {rw.get('reason', '')}")

    if rewrites or opt_summary or suggested_skills:
        if st.button("⬇️ Apply All Rewrites to Editor", type="primary"):
            base = st.session_state.get("resume_editor", st.session_state.resume_text)
            not_found = []
            current_summary = a.get("current_summary", "")
            if current_summary and "summary_edit" in st.session_state:
                base, ok = _fuzzy_replace(base, current_summary, st.session_state["summary_edit"])
                if not ok:
                    not_found.append("Summary")
            current_skills = a.get("current_skills", "")
            if current_skills and "skills_edit" in st.session_state:
                base, ok = _fuzzy_replace(base, current_skills, st.session_state["skills_edit"])
                if not ok:
                    not_found.append("Skills")
            for i, rw in enumerate(rewrites):
                key = f"rw_edit_{i}"
                orig = rw.get("original", "")
                if orig and key in st.session_state:
                    base, ok = _fuzzy_replace(base, orig, st.session_state[key])
                    if not ok:
                        not_found.append(f"Rewrite #{i + 1}")
            st.session_state["resume_editor"] = base
            if not_found:
                st.warning(
                    f"Could not locate: **{', '.join(not_found)}**. "
                    "Apply those manually in the Edit Source tab."
                )
            else:
                st.success("All rewrites applied! See the updated preview on the right →")

        if st.button("🪄 Auto-Format Resume", type="secondary",
                     help="Use AI to restructure the resume for better ATS formatting score"):
            base = st.session_state.get("resume_editor", st.session_state.resume_text)
            with st.spinner("Reformatting…"):
                try:
                    st.session_state["resume_editor"] = auto_format_resume(base)
                    st.success("Resume reformatted! Check the Preview tab.")
                except RuntimeError as exc:
                    st.error(str(exc))
                except anthropic.APIError as exc:
                    st.error(f"Anthropic API error: {exc}")

    st.divider()

    # 5 ── Two columns: Keyword Gaps + Priority  |  Preview + Edit ────────────────
    left_col, right_col = st.columns([1, 1])

    with left_col:
        st.subheader("🔍 Keyword Gaps")
        gaps = a.get("keyword_gaps", {})
        category_labels = {
            "technical_skills":   "🔧 Technical Skills",
            "soft_skills":        "🤝 Soft Skills",
            "tools_technologies": "🛠️ Tools & Technologies",
            "certifications":     "🏆 Certifications",
            "domain_knowledge":   "📚 Domain Knowledge",
        }
        any_gaps = False
        for key, label in category_labels.items():
            items = gaps.get(key, [])
            if items:
                any_gaps = True
                with st.expander(f"{label} — {len(items)} missing", expanded=len(items) <= 6):
                    st.markdown("\n".join(f"- `{item}`" for item in items))
        if not any_gaps:
            st.success("No significant keyword gaps detected!")

    with right_col:
        tab_preview, tab_edit = st.tabs(["👁️ Preview", "✏️ Edit Source"])
        with tab_preview:
            preview_text = st.session_state.get("resume_editor", st.session_state.resume_text)
            st.markdown(format_resume_html(preview_text), unsafe_allow_html=True)
        with tab_edit:
            st.caption("Use `**word**` to bold text. Changes reflect in Preview and the download.")
            st.text_area(
                "Resume editor",
                value=st.session_state.resume_text,
                height=520,
                key="resume_editor",
                label_visibility="collapsed",
            )

    st.divider()

    # 6 ── Score edited resume ─────────────────────────────────────────────────────
    rescore_jd = st.session_state.get("analysis_jd", "")
    if rescore_jd:
        if st.button("📊 Score Edited Resume", type="secondary"):
            export_text = st.session_state.get("resume_editor", st.session_state.resume_text)
            with st.spinner("Scoring edited resume…"):
                try:
                    st.session_state.rescore = analyze_resume(export_text, rescore_jd)
                except Exception as exc:
                    st.error(f"Scoring failed: {exc}")

    if st.session_state.get("rescore"):
        rs = st.session_state.rescore
        orig = st.session_state.analysis
        rs_bk = rs.get("score_breakdown", {})
        orig_bk = orig.get("score_breakdown", {})

        def _delta(key):
            try:
                return int(rs_bk.get(key, 0)) - int(orig_bk.get(key, 0))
            except (TypeError, ValueError):
                return None

        new_score = int(rs.get("ats_score_estimate", 0))
        old_score = int(orig.get("ats_score_estimate", 0))
        st.subheader("📈 Updated Score")
        rm1, rm2, rm3, rm4, rm5 = st.columns(5)
        rm1.metric("Overall ATS Score", f"{new_score} / 100", delta=new_score - old_score)
        rm2.metric("Keyword Match", rs_bk.get("keyword_match", "—"), delta=_delta("keyword_match"))
        rm3.metric("Quantification", rs_bk.get("quantification", "—"), delta=_delta("quantification"))
        rm4.metric("Formatting", rs_bk.get("formatting", "—"), delta=_delta("formatting"))
        rm5.metric("Relevance", rs_bk.get("relevance", "—"), delta=_delta("relevance"))
        st.progress(min(new_score, 100) / 100)

# ── Resume download ───────────────────────────────────────────────────────────────
if st.session_state.resume_text:
    export_text = st.session_state.get("resume_editor", st.session_state.resume_text)
    st.download_button(
        label="📥 Download Resume as Word (.docx)",
        data=to_word(export_text),
        file_name=st.session_state.export_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# ── Cover letter ──────────────────────────────────────────────────────────────────
if st.session_state.resume_text:
    st.divider()
    st.subheader("✉️ Cover Letter")

    cl_jd = st.session_state.get("analysis_jd", "")
    if not cl_jd:
        st.info("Run the resume analysis first to enable cover letter generation.")
    else:
        if st.button("✍️ Generate Cover Letter", type="primary"):
            with st.spinner("Writing your cover letter…"):
                try:
                    letter = generate_cover_letter(st.session_state.resume_text, cl_jd)
                    st.session_state.cover_letter = letter
                    if "cl_editor" in st.session_state:
                        del st.session_state["cl_editor"]
                except RuntimeError as exc:
                    st.error(str(exc))
                except anthropic.APIError as exc:
                    st.error(f"Anthropic API error: {exc}")

        if st.session_state.cover_letter:
            tab_cl_prev, tab_cl_edit = st.tabs(["👁️ Preview", "✏️ Edit"])
            with tab_cl_prev:
                cl_display = st.session_state.get("cl_editor", st.session_state.cover_letter)
                st.markdown(format_letter_html(cl_display), unsafe_allow_html=True)
            with tab_cl_edit:
                st.text_area(
                    "Cover letter editor",
                    value=st.session_state.cover_letter,
                    height=420,
                    key="cl_editor",
                    label_visibility="collapsed",
                )

            cl_export = st.session_state.get("cl_editor", st.session_state.cover_letter)
            cl_filename = st.session_state.export_filename.replace(
                "_optimized.docx", "_cover_letter.docx"
            )
            st.download_button(
                label="📥 Download Cover Letter as Word (.docx)",
                data=to_word_letter(cl_export),
                file_name=cl_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
